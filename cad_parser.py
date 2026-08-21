"""
cad_parser.py - CAD 解析與報表產出模組
Version: v2.6.0_20260821
Description: Phase 1 工程三視圖與黑白 ISO View 組合圖卡產出。
             採用第三角法 (Third-Angle Projection)：TOP, FRONT, RIGHT 與 ISOMETRIC。
             透過 3D BoundingBox 計算全域統一像素比例 (Global Co-Scale Mechanism)，
             確保三視圖長寬比例一致不變形，並自動執行跨視圖投影對齊 (Projection Alignment)。
             包含獨立 UUID 檔名隔離、Pillow 圖片有效性驗證與 image_error 診斷回傳。
             獨立匯出 Word 圖文報價單 (.docx)，100% 不干擾 Excel 原有導出邏輯與公式運算。
"""

import os
import math
import uuid
import tempfile
from datetime import datetime
from typing import Dict, Any, List, Optional
import cadquery as cq
from PIL import Image, ImageDraw, ImageFont
import openpyxl
from openpyxl.styles import Font

# 引入 docx 模組
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 引入 CairoSVG 向量轉譯模組
try:
    import cairosvg
    HAS_CAIROSVG = True
except ImportError:
    HAS_CAIROSVG = False

# 引入 OpenCASCADE 原生 Bnd_OBB 模組以進行精確幾何最小包容盒計算
try:
    from OCP.Bnd import Bnd_OBB
    from OCP.BRepBndLib import BRepBndLib
    HAS_OCP_OBB = True
except ImportError:
    HAS_OCP_OBB = False


def safe_str(val: Any) -> str:
    """安全轉為字串並去除首尾空白"""
    if val is None:
        return ""
    try:
        return str(val).strip()
    except Exception:
        return ""


def set_cell_value_safe(ws, row: int, col: int, value: Any, font: Optional[Font] = None):
    """安全填寫儲存格：若遇到 MergedCell，寫入其最左上角的 Master Cell"""
    try:
        cell = ws.cell(row=row, column=col)
        if type(cell).__name__ == 'MergedCell':
            for rng in ws.merged_cells.ranges:
                if cell.coordinate in rng:
                    master_cell = ws.cell(row=rng.min_row, column=rng.min_col)
                    master_cell.value = value
                    if font:
                        master_cell.font = font
                    return
        cell.value = value
        if font:
            cell.font = font
    except Exception:
        pass


def is_valid_image(path: Optional[str]) -> bool:
    """驗證圖片檔是否存在、大小大於 0、且可被 Pillow 正確讀取"""
    if not path or not os.path.exists(path):
        return False
    if os.path.getsize(path) <= 0:
        return False
    try:
        with Image.open(path) as img:
            img.verify()
        return True
    except Exception:
        return False


def calculate_obb_dimensions(model: cq.Workplane) -> List[float]:
    """使用 OpenCASCADE 原生 Bnd_OBB 精確計算 3D 實體的最小素材包容盒 (OBB)"""
    if HAS_OCP_OBB:
        try:
            occ_shape = model.val().wrapped
            obb = Bnd_OBB()
            BRepBndLib.AddOBB_s(occ_shape, obb, True, True, True)
            dims = [obb.XHSize() * 2.0, obb.YHSize() * 2.0, obb.ZHSize() * 2.0]
            if all(d > 0 for d in dims):
                return dims
        except Exception:
            pass

    bbox = model.val().BoundingBox()
    return [bbox.xlen, bbox.ylen, bbox.zlen]


def render_single_view(
    model: cq.Workplane, 
    proj_dir: tuple, 
    out_png_path: str, 
    out_w: int = 600, 
    out_h: int = 600
) -> bool:
    """將模型以指定的投影方向導出為 SVG，並使用 cairosvg 轉為指定大小之 PNG"""
    temp_dir = tempfile.gettempdir()
    unique_id = str(uuid.uuid4())[:8]
    svg_path = os.path.join(temp_dir, f"view_{unique_id}.svg")

    try:
        cq.exporters.export(model, svg_path, opt={"projectionDir": proj_dir, "showHidden": False})
        if not os.path.exists(svg_path) or os.path.getsize(svg_path) <= 0:
            return False

        if HAS_CAIROSVG:
            cairosvg.svg2png(url=svg_path, write_to=out_png_path, output_width=out_w, output_height=out_h)
            return is_valid_image(out_png_path)
        return False
    except Exception:
        return False
    finally:
        if os.path.exists(svg_path):
            try:
                os.remove(svg_path)
            except Exception:
                pass


def auto_crop_image_to_bbox(input_path: str, margin_px: int = 10) -> Optional[Image.Image]:
    """精確裁切 PNG 圖像中非背景區域，並附加安全像素 Margin"""
    if not is_valid_image(input_path):
        return None
    try:
        with Image.open(input_path) as img:
            img_rgba = img.convert("RGBA")
            w, h = img_rgba.size
            mask = Image.new("L", (w, h), 0)
            pixels_rgba = img_rgba.load()
            pixels_mask = mask.load()

            for y in range(h):
                for x in range(w):
                    r, g, b, a = pixels_rgba[x, y]
                    if a > 10 and (r < 240 or g < 240 or b < 240):
                        pixels_mask[x, y] = 255

            bbox = mask.getbbox()
            if not bbox:
                return img_rgba.copy()

            l, u, r, d = bbox
            nl = max(0, l - margin_px)
            nu = max(0, u - margin_px)
            nr = min(w, r + margin_px)
            nd = min(h, d + margin_px)
            return img_rgba.crop((nl, nu, nr, nd))
    except Exception:
        return None


def generate_engineering_4views_card(model: cq.Workplane, final_png_path: str) -> bool:
    """
     Phase 1 核心：產生第三角法三視圖 (TOP, FRONT, RIGHT) 與黑白 ISO 組合圖卡。
    落實共同 Scale Factor 與投影軸線 (Projection Alignment) 對齊。
    """
    bbox_3d = model.val().BoundingBox()
    x_len = max(0.1, bbox_3d.xlen)
    y_len = max(0.1, bbox_3d.ylen)
    z_len = max(0.1, bbox_3d.zlen)

    # 1. 計算三視圖跨距比例 (Global Scale Factor)
    max_span = max(x_len, y_len, z_len)
    cell_draw_area = 550.0  # 單區塊內部最大像素繪圖區域

    # 計算各 View 的統一物理對應像素大小
    scale_ratio = cell_draw_area / max_span
    
    # TOP View (X x Y)
    top_w = max(40, int(x_len * scale_ratio))
    top_h = max(40, int(y_len * scale_ratio))
    
    # FRONT View (X x Z)
    front_w = max(40, int(x_len * scale_ratio))
    front_h = max(40, int(z_len * scale_ratio))
    
    # RIGHT View (Y x Z)
    right_w = max(40, int(y_len * scale_ratio))
    right_h = max(40, int(z_len * scale_ratio))

    temp_dir = tempfile.gettempdir()
    uid = str(uuid.uuid4())[:8]
    p_top = os.path.join(temp_dir, f"top_{uid}.png")
    p_front = os.path.join(temp_dir, f"front_{uid}.png")
    p_right = os.path.join(temp_dir, f"right_{uid}.png")
    p_iso = os.path.join(temp_dir, f"iso_{uid}.png")

    try:
        # 2. 依據投影方向匯出 PNG
        # TOP: (0, 0, 1) | FRONT: (0, -1, 0) | RIGHT: (1, 0, 0) | ISO: (1, 1, 1)
        ok_t = render_single_view(model, (0, 0, 1), p_top, top_w + 100, top_h + 100)
        ok_f = render_single_view(model, (0, -1, 0), p_front, front_w + 100, front_h + 100)
        ok_r = render_single_view(model, (1, 0, 0), p_right, right_w + 100, right_h + 100)
        ok_i = render_single_view(model, (1, 1, 1), p_iso, 600, 600)

        if not (ok_t and ok_f and ok_r and ok_i):
            return False

        # 3. 讀取並裁剪邊界 (Crop)
        img_top = auto_crop_image_to_bbox(p_top, margin_px=6)
        img_front = auto_crop_image_to_bbox(p_front, margin_px=6)
        img_right = auto_crop_image_to_bbox(p_right, margin_px=6)
        img_iso = auto_crop_image_to_bbox(p_iso, margin_px=12)

        if not (img_top and img_front and img_right and img_iso):
            return False

        # 4. 建立 1600x1200 畫布並進行第三角法對齊組合
        canvas_w, canvas_h = 1600, 1200
        canvas = Image.new("RGBA", (canvas_w, canvas_h), (255, 255, 255, 255))

        # 定義 2x2 四個分區中心點
        # 左上: TOP (400, 300) | 右上: ISO (1200, 300)
        # 左下: FRONT (400, 900) | 右下: RIGHT (1200, 900)
        cx_left, cx_right = 420, 1180
        cy_top, cy_bottom = 300, 900

        # A. 放置 FRONT View (左下錨點)
        fw, fh = img_front.size
        front_x = cx_left - (fw // 2)
        front_y = cy_bottom - (fh // 2)
        canvas.paste(img_front, (front_x, front_y), img_front)

        # B. 放置 TOP View (左上，強制作圖 X 軸水平對齊)
        tw, th = img_top.size
        top_x = front_x + (fw // 2) - (tw // 2)  # 水平對齊 FRONT 中心
        top_y = cy_top - (th // 2)
        canvas.paste(img_top, (top_x, top_y), img_top)

        # C. 放置 RIGHT View (右下，強制作圖 Y 軸垂直對齊)
        rw, rh = img_right.size
        right_x = cx_right - (rw // 2)
        right_y = front_y + (fh // 2) - (rh // 2)  # 垂直對齊 FRONT 中心
        canvas.paste(img_right, (right_x, right_y), img_right)

        # D. 放置 ISO View (右上，獨立放大為視覺參考)
        iw, ih = img_iso.size
        iso_x = cx_right - (iw // 2)
        iso_y = cy_top - (ih // 2)
        canvas.paste(img_iso, (iso_x, iso_y), img_iso)

        # 5. 繪製標籤與第三角法標準標註
        draw = ImageDraw.Draw(canvas)
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 22)
            font_small = ImageFont.truetype("DejaVuSans.ttf", 16)
        except Exception:
            font = ImageFont.load_default()
            font_small = font

        # 左上角標註標準
        draw.text((30, 25), "PROJECTION: THIRD ANGLE", fill=(100, 100, 100, 255), font=font_small)

        # 各分區標籤
        draw.text((cx_left - 60, cy_top + 230), "平面圖 TOP", fill=(60, 60, 60, 255), font=font)
        draw.text((cx_left - 65, cy_bottom + 230), "正面圖 FRONT", fill=(60, 60, 60, 255), font=font)
        draw.text((cx_right - 65, cy_bottom + 230), "右視圖 RIGHT", fill=(60, 60, 60, 255), font=font)
        draw.text((cx_right - 80, cy_top + 230), "等角圖 ISOMETRIC", fill=(60, 60, 60, 255), font=font)

        # 畫布中央輕微十字分隔線
        draw.line([(800, 80), (800, 1120)], fill=(230, 230, 230, 255), width=2)
        draw.line([(80, 600), (1520, 600)], fill=(230, 230, 230, 255), width=2)

        canvas.convert("RGB").save(final_png_path, "PNG")
        return is_valid_image(final_png_path)

    except Exception as e_card:
        print(f"[ENGINEERING CARD ERROR] {str(e_card)}")
        return False
    finally:
        for p in [p_top, p_front, p_right, p_iso]:
            if os.path.exists(p):
                try:
                    os.remove(p)
                except Exception:
                    pass


def parse_cad_with_screenshot(file_path: str) -> Dict[str, Any]:
    """
    讀取 CAD 檔案，同時計算 OBB 與 AABB 尺寸，自動採納素材體積較小者。
    呼叫 generate_engineering_4views_card 產出三視圖圖卡。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"找不到檔案：{file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    valid_extensions = ('.step', '.stp', '.igs', '.iges')
    if ext not in valid_extensions:
        raise ValueError(f"不支援的格式 '{ext}'")

    # 1. 幾何長寬高雙軌解析與體積最小化比對
    try:
        if ext in ('.step', '.stp'):
            model = cq.importers.importStep(file_path)
        elif ext in ('.igs', '.iges'):
            try:
                model = cq.importers.importShape(file_path)
            except Exception:
                from OCP.IGESControl import IGESControl_Reader
                reader = IGESControl_Reader()
                reader.ReadFile(file_path)
                reader.TransferRoots()
                occ_shape = reader.Shape()
                model = cq.Workplane("XY").newObject([cq.Shape.cast(occ_shape)])

        # A. OBB 最小包容盒運算
        raw_dims_obb = calculate_obb_dimensions(model)
        x_obb = math.ceil(raw_dims_obb[0])
        y_obb = math.ceil(raw_dims_obb[1])
        z_obb = math.ceil(raw_dims_obb[2])
        sorted_obb = sorted([x_obb, y_obb, z_obb], reverse=True)
        vol_obb = sorted_obb[0] * sorted_obb[1] * sorted_obb[2]

        # B. AABB 標準投影外框運算
        bbox = model.val().BoundingBox()
        x_aabb = math.ceil(bbox.xlen)
        y_aabb = math.ceil(bbox.ylen)
        z_aabb = math.ceil(bbox.zlen)
        sorted_aabb = sorted([x_aabb, y_aabb, z_aabb], reverse=True)
        vol_aabb = sorted_aabb[0] * sorted_aabb[1] * sorted_aabb[2]

        # C. 自動採納體積較小者 (Min-Volume Selection)
        if vol_obb <= vol_aabb:
            length_val, width_val, height_val = sorted_obb[0], sorted_obb[1], sorted_obb[2]
            used_mode = "OBB"
        else:
            length_val, width_val, height_val = sorted_aabb[0], sorted_aabb[1], sorted_aabb[2]
            used_mode = "AABB"

        dimensions_str = f"{length_val}*{width_val}*{height_val}"
    except Exception as e:
        return {
            "status": "error",
            "file_name": os.path.basename(file_path),
            "error_message": f"CAD 幾何解析失敗: {str(e)}"
        }

    # 2. 獨立第三角法工程視圖圖卡繪製 (UUID 檔名隔離，捕捉完整 image_error)
    img_path = None
    image_error = None
    
    unique_id = str(uuid.uuid4())[:8]
    temp_dir = tempfile.gettempdir()
    final_card_png = os.path.join(temp_dir, f"cad_card_{unique_id}.png")

    try:
        success = generate_engineering_4views_card(model, final_card_png)
        if success and is_valid_image(final_card_png):
            img_path = final_card_png
        else:
            raise ValueError("工程三視圖圖卡合成失敗或產出無效圖片。")

    except Exception as e_img:
        image_error = f"{type(e_img).__name__}: {str(e_img)}"
        print(f"[CAD IMAGE ERROR] {os.path.basename(file_path)}: {image_error}")
        img_path = None

    return {
        "status": "success",
        "file_name": os.path.basename(file_path),
        "dimensions_str": dimensions_str,
        "length": length_val,
        "width": width_val,
        "height": height_val,
        "unit": "mm",
        "used_mode": used_mode,
        "image_path": img_path,
        "image_error": image_error
    }


def generate_excel_report(
    parsed_results: List[Dict[str, Any]], 
    output_excel_path: str,
    header_info: Optional[Dict[str, Any]] = None
):
    """載入範本檔寫入數據，全數指定字體為標楷體，自動調整欄寬"""
    template_candidates = [
        "template.xlsm", "template.xlsx", "template.xls",
        "Template.xlsm", "Template.xlsx", "Template.xls"
    ]
    template_file = None
    for tf in template_candidates:
        if os.path.exists(tf):
            template_file = tf
            break

    is_xlsm = template_file and template_file.lower().endswith('.xlsm')

    if template_file:
        wb = openpyxl.load_workbook(template_file, data_only=False, keep_vba=is_xlsm)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "報價單"

    kai_font_regular = Font(name="標楷體", size=11, bold=False)
    kai_font_bold = Font(name="標楷體", size=11, bold=True)

    if header_info and isinstance(header_info, dict):
        customer_val = safe_str(header_info.get("customer"))
        contact_val = safe_str(header_info.get("contact"))
        phone_val = safe_str(header_info.get("phone"))
        fax_val = safe_str(header_info.get("fax"))

        if customer_val:
            set_cell_value_safe(ws, 4, 2, f"客戶名稱 : {customer_val}", font=kai_font_regular)
        if contact_val:
            set_cell_value_safe(ws, 5, 2, f"聯 絡 人 : {contact_val}", font=kai_font_regular)
        if phone_val:
            set_cell_value_safe(ws, 6, 2, f"聯絡電話 : {phone_val}", font=kai_font_regular)
        if fax_val:
            set_cell_value_safe(ws, 7, 2, f"傳    真 : {fax_val}", font=kai_font_regular)

    today_str = datetime.now().strftime("%Y/%m/%d")
    set_cell_value_safe(ws, 7, 15, today_str, font=kai_font_regular)

    start_row = 10
    max_b_len = 12
    max_p_len = 16

    for idx, item in enumerate(parsed_results):
        row_num = start_row + idx
        
        cell_a = ws.cell(row=row_num, column=1, value=idx + 1)
        cell_a.font = kai_font_regular
        
        file_name = item.get("file_name", "")
        cell_b = ws.cell(row=row_num, column=2, value=file_name)
        cell_b.font = kai_font_regular
        max_b_len = max(max_b_len, len(str(file_name)))
        
        dims_str = item.get("dimensions_str", "")
        cell_p = ws.cell(row=row_num, column=16, value=dims_str)
        cell_p.font = kai_font_bold
        max_p_len = max(max_p_len, len(str(dims_str)))
        
        if "length" in item:
            cell_q = ws.cell(row=row_num, column=17, value=item["length"])
            cell_q.font = kai_font_bold
        if "width" in item:
            cell_r = ws.cell(row=row_num, column=18, value=item["width"])
            cell_r.font = kai_font_bold
        if "height" in item:
            cell_s = ws.cell(row=row_num, column=19, value=item["height"])
            cell_s.font = kai_font_bold
            
        cell_t = ws.cell(row=row_num, column=20, value=item.get("unit", "mm"))
        cell_t.font = kai_font_regular

        c_g = ws.cell(row=row_num, column=7, value=f"=E{row_num}*F{row_num}")
        c_j = ws.cell(row=row_num, column=10, value=f"=H{row_num}*I{row_num}")
        c_m = ws.cell(row=row_num, column=13, value=f"=K{row_num}*L{row_num}")
        c_g.font = kai_font_regular
        c_j.font = kai_font_regular
        c_m.font = kai_font_regular

    ws.column_dimensions['B'].width = max(ws.column_dimensions['B'].width or 0, max_b_len + 6)
    ws.column_dimensions['P'].width = max(ws.column_dimensions['P'].width or 0, max_p_len + 6)

    total_row = 101
    for r in range(10, ws.max_row + 1):
        cell_a_val = str(ws.cell(row=r, column=1).value or "").replace(" ", "")
        cell_b_val = str(ws.cell(row=r, column=2).value or "").replace(" ", "")
        cell_c_val = str(ws.cell(row=r, column=3).value or "").replace(" ", "")
        
        if "合計" in cell_a_val or "合計" in cell_b_val or "合計" in cell_c_val:
            total_row = r
            break

    data_end_row = total_row - 1
    set_cell_value_safe(ws, total_row, 7, f"=SUM(G10:G{data_end_row})", font=kai_font_bold)
    set_cell_value_safe(ws, total_row, 10, f"=SUM(J10:J{data_end_row})", font=kai_font_bold)
    set_cell_value_safe(ws, total_row, 13, f"=SUM(M10:M{data_end_row})", font=kai_font_bold)
    set_cell_value_safe(ws, total_row, 15, f"=G{total_row}+J{total_row}+M{total_row}", font=kai_font_bold)

    wb.save(output_excel_path)


def generate_word_report(
    parsed_results: List[Dict[str, Any]], 
    output_word_path: str,
    header_info: Optional[Dict[str, Any]] = None
):
    """建立 Word 圖文報價單 (.docx)，插入工程三視圖圖卡 (寬度 5.8 英吋)"""
    if not HAS_DOCX:
        raise ModuleNotFoundError("系統缺少 python-docx 套件，無法產生 Word 報表。")

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("⚙️ CAD 零件報價與工程三視圖明細單")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.name = '標楷體'

    today_str = datetime.now().strftime("%Y年%m月%d日")
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub_run = sub_p.add_run(f"報價日期：{today_str}")
    sub_run.font.size = Pt(10)
    sub_run.font.name = '標楷體'

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if header_info and isinstance(header_info, dict):
        cust = safe_str(header_info.get("customer")) or "未填寫"
        contact = safe_str(header_info.get("contact")) or "未填寫"
        phone = safe_str(header_info.get("phone")) or "未填寫"
        fax = safe_str(header_info.get("fax")) or "未填寫"

        info_table = doc.add_table(rows=2, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.autofit = False

        cells = info_table.rows[0].cells
        cells[0].text = f"客戶名稱：{cust}"
        cells[1].text = f"聯 絡 人：{contact}"

        cells_2 = info_table.rows[1].cells
        cells_2[0].text = f"聯絡電話：{phone}"
        cells_2[1].text = f"傳    真：{fax}"

        for row in info_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '標楷體'
                        run.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    for idx, item in enumerate(parsed_results):
        if item.get("status") != "success":
            continue

        file_name = item.get("file_name", "")
        dims_str = item.get("dimensions_str", "")
        used_mode = item.get("used_mode", "OBB")
        length = item.get("length", 0)
        width = item.get("width", 0)
        height = item.get("height", 0)
        unit = item.get("unit", "mm")
        img_path = item.get("image_path")
        img_err = item.get("image_error")

        p_item = doc.add_paragraph()
        r_item = p_item.add_run(f"【項目 {idx+1}】 檔名：{file_name}")
        r_item.font.bold = True
        r_item.font.size = Pt(12)
        r_item.font.name = '標楷體'

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.2)
        
        r_d1 = p_desc.add_run(f"• 採納素材尺寸 (長*寬*高)：{dims_str} {unit} ({used_mode} 模式)\n")
        r_d1.font.bold = True
        r_d1.font.name = '標楷體'
        
        r_d2 = p_desc.add_run(f"• 尺寸拆解數值：長 {length} {unit} / 寬 {width} {unit} / 高 {height} {unit}\n")
        r_d2.font.name = '標楷體'

        # 插入工程三視圖圖卡 (寬度設為 Inches(5.8) 滿版展現)
        if is_valid_image(img_path):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Inches(5.8))
            except Exception as e_word_img:
                p_err = doc.add_paragraph(f"   [ 無法產生工程視圖預覽 (Word插入例外: {str(e_word_img)}) ]")
                p_err.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        else:
            err_msg = f" ({img_err})" if img_err else ""
            p_none = doc.add_paragraph(f"   [ 無法產生工程視圖預覽{err_msg} ]")
            p_none.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.save(output_word_path)
